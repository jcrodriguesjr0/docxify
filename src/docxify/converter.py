"""Core MD -> DOCX conversion engine."""

from __future__ import annotations

import os
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .styles import StyleProfile, PRESETS
from .markdown_parser import (
    parse_markdown,
    Node, Heading, Paragraph, BulletItem, NumberedItem,
    CodeBlock, BlockQuote, HorizontalRule, BlankLine,
    TextRun, RunStyle,
)


class DocxBuilder:
    """Builds a .docx document from parsed markdown nodes."""

    def __init__(self, style: StyleProfile | None = None, template: str | None = None):
        self.style = style or PRESETS["default"]

        if template and os.path.exists(template):
            self.doc = Document(template)
            self._clear_body()
        else:
            self.doc = Document()

        self._setup_page()
        self._setup_default_style()

    def _clear_body(self) -> None:
        body = self.doc.element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)

    def _setup_page(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Cm(self.style.page_width_cm)
        section.page_height = Cm(self.style.page_height_cm)
        section.top_margin = Cm(self.style.margin_top_cm)
        section.bottom_margin = Cm(self.style.margin_bottom_cm)
        section.left_margin = Cm(self.style.margin_left_cm)
        section.right_margin = Cm(self.style.margin_right_cm)

    def _setup_default_style(self) -> None:
        style = self.doc.styles["Normal"]
        style.font.name = self.style.font_name
        style.font.size = Pt(self.style.body_size)
        style.font.color.rgb = RGBColor.from_string(self.style.body_color)
        style.paragraph_format.space_before = Pt(self.style.space_before)
        style.paragraph_format.space_after = Pt(self.style.space_after)

        # Language
        rPr = style.element.get_or_add_rPr()
        lang = rPr.makeelement(qn("w:lang"), {
            qn("w:val"): self.style.language,
            qn("w:eastAsia"): self.style.language,
        })
        rPr.append(lang)

    def _add_runs(self, paragraph, runs: list[TextRun]) -> None:
        for run_data in runs:
            run = paragraph.add_run(run_data.text)
            run.font.name = self.style.font_name
            run.font.size = Pt(self.style.body_size)
            run.font.color.rgb = RGBColor.from_string(self.style.body_color)

            if run_data.style == RunStyle.BOLD:
                run.bold = True
            elif run_data.style == RunStyle.ITALIC:
                run.italic = True
            elif run_data.style == RunStyle.BOLD_ITALIC:
                run.bold = True
                run.italic = True
            elif run_data.style == RunStyle.CODE:
                run.font.name = self.style.code_font
                run.font.size = Pt(self.style.code_size)

    def _add_heading(self, node: Heading) -> None:
        p = self.doc.add_paragraph()
        align_str = self.style.heading_align(node.level)
        p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align_str == "center"
                       else WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_before = Pt(self.style.space_before * 2)
        p.paragraph_format.space_after = Pt(self.style.space_after)

        size = self.style.heading_size(node.level)
        for run_data in node.runs:
            run = p.add_run(run_data.text)
            run.font.name = self.style.font_name
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor.from_string(self.style.heading_color)
            run.bold = True

    def _add_paragraph(self, node: Paragraph) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._add_runs(p, node.runs)

    def _add_bullet(self, node: BulletItem) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        indent_cm = 1.27 * (node.indent_level + 1)
        p.paragraph_format.left_indent = Cm(indent_cm)

        bullet_run = p.add_run("\u2022 ")
        bullet_run.font.name = self.style.font_name
        bullet_run.font.size = Pt(self.style.body_size)

        self._add_runs(p, node.runs)

    def _add_numbered(self, node: NumberedItem) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        indent_cm = 1.27 * (node.indent_level + 1)
        p.paragraph_format.left_indent = Cm(indent_cm)

        num_run = p.add_run(f"{node.number}. ")
        num_run.font.name = self.style.font_name
        num_run.font.size = Pt(self.style.body_size)

        self._add_runs(p, node.runs)

    def _add_code_block(self, node: CodeBlock) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(self.style.space_before)
        p.paragraph_format.space_after = Pt(self.style.space_after)

        # Background shading
        pPr = p._element.get_or_add_pPr()
        shd = pPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): self.style.code_bg_color,
        })
        pPr.append(shd)

        run = p.add_run(node.code)
        run.font.name = self.style.code_font
        run.font.size = Pt(self.style.code_size)

    def _add_blockquote(self, node: BlockQuote) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(1.0)

        # Left border via XML
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        left = pBdr.makeelement(qn("w:left"), {
            qn("w:val"): "single",
            qn("w:sz"): "12",
            qn("w:space"): "4",
            qn("w:color"): "CCCCCC",
        })
        pBdr.append(left)
        pPr.append(pBdr)

        self._add_runs(p, node.runs)

    def _add_hr(self) -> None:
        p = self.doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): "999999",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_blank(self) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    def render(self, nodes: list[Node]) -> None:
        for node in nodes:
            if isinstance(node, Heading):
                self._add_heading(node)
            elif isinstance(node, Paragraph):
                self._add_paragraph(node)
            elif isinstance(node, BulletItem):
                self._add_bullet(node)
            elif isinstance(node, NumberedItem):
                self._add_numbered(node)
            elif isinstance(node, CodeBlock):
                self._add_code_block(node)
            elif isinstance(node, BlockQuote):
                self._add_blockquote(node)
            elif isinstance(node, HorizontalRule):
                self._add_hr()
            elif isinstance(node, BlankLine):
                self._add_blank()

    def save(self, output_path: str | Path) -> Path:
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))
        return path


def convert(
    markdown_text: str,
    output_path: str | Path,
    style: StyleProfile | None = None,
    template: str | None = None,
) -> Path:
    """Convert a markdown string to a .docx file.

    Args:
        markdown_text: Raw markdown content.
        output_path: Where to save the .docx file.
        style: Styling profile. Uses default if None.
        template: Optional path to a .docx template file.

    Returns:
        Path to the generated .docx file.
    """
    nodes = parse_markdown(markdown_text)
    builder = DocxBuilder(style=style, template=template)
    builder.render(nodes)
    return builder.save(output_path)
