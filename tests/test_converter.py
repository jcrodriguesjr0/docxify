"""Tests for MD -> DOCX conversion."""

from pathlib import Path
from docx import Document

from docxify.converter import convert
from docxify.styles import StyleProfile, PRESETS


def test_basic_conversion(sample_md, tmp_docx):
    result = convert(sample_md, tmp_docx)
    assert result.exists()
    assert result.suffix == ".docx"

    doc = Document(str(result))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("Project Report" in t for t in texts)
    assert any("Executive Summary" in t for t in texts)


def test_headings_preserved(sample_md, tmp_docx):
    convert(sample_md, tmp_docx)
    doc = Document(str(tmp_docx))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "Project Report" in texts
    assert "Executive Summary" in texts
    assert "Key Findings" in texts


def test_bullets_rendered(sample_md, tmp_docx):
    convert(sample_md, tmp_docx)
    doc = Document(str(tmp_docx))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "First bullet point" in full_text
    assert "Second bullet" in full_text


def test_code_block_rendered(sample_md, tmp_docx):
    convert(sample_md, tmp_docx)
    doc = Document(str(tmp_docx))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "def hello():" in full_text


def test_preset_styles(sample_md, tmp_docx):
    for name in PRESETS:
        convert(sample_md, tmp_docx, style=PRESETS[name])
        assert tmp_docx.exists()


def test_custom_style(sample_md, tmp_docx):
    style = StyleProfile(font_name="Arial", body_size=12, language="pt-BR")
    convert(sample_md, tmp_docx, style=style)

    doc = Document(str(tmp_docx))
    normal_style = doc.styles["Normal"]
    assert normal_style.font.name == "Arial"


def test_creates_parent_dirs(sample_md, tmp_path):
    deep_path = tmp_path / "a" / "b" / "c" / "output.docx"
    convert(sample_md, deep_path)
    assert deep_path.exists()


def test_empty_markdown(tmp_docx):
    convert("", tmp_docx)
    assert tmp_docx.exists()
